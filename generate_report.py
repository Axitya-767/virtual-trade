from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# --- Style Setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_section_heading(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()

def add_subheading(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)

def add_sub3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.size = Pt(12)

def add_body(text):
    doc.add_paragraph(text)

def add_bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    # Light gray background via shading
    shading = qn('w:shd')
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(shading, {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F2F2F2'
    })
    pPr.append(shd)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def add_query_block(number, title, sql, purpose):
    p = doc.add_paragraph()
    run = p.add_run(f"Query {number}: {title}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    add_code(sql)
    p2 = doc.add_paragraph()
    run2 = p2.add_run("Purpose: ")
    run2.font.bold = True
    run2.font.size = Pt(10)
    run3 = p2.add_run(purpose)
    run3.font.size = Pt(10)
    doc.add_paragraph()


# ==========================================
# COVER PAGE
# ==========================================
for _ in range(6):
    doc.add_paragraph()

add_title("DBMS Project Report")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Virtual Trade Engine")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 102, 51)
run.font.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("A Real-Time Stock & Crypto Trading Simulator")
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(100, 100, 100)

for _ in range(4):
    doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Web Programming & DBMS Project")
run3.font.size = Pt(12)

doc.add_page_break()

# ==========================================
# I. STORYLINE
# ==========================================
add_section_heading("I. Storyline")

add_body("Problem Statement: Real-world stock trading platforms are complex and risky for beginners. New investors often lose money because they lack a safe environment to learn how markets work, how prices fluctuate, and how buy/sell decisions affect their portfolio.")

add_body("Our Solution: We built VirtualTrade — a full-stack web application that simulates a live stock market. Every user starts with a virtual wallet of $10,000 and can:")

add_bullet("View 5 real-world assets (Apple, Tesla, Bitcoin, Ethereum, Google) with prices that change every 3 seconds")
add_bullet("Buy and sell assets using virtual money")
add_bullet("Track their portfolio (what they own and its current value)")
add_bullet("View a full transaction history of all their trades")

add_body("Why DBMS matters here: Every trade involves multiple database operations that must happen atomically — deducting money from the wallet, adding the asset to the portfolio, and recording the transaction. If any step fails, the entire trade must roll back. This is a textbook use case for ACID transactions, which is the core DBMS concept demonstrated in this project.")

add_body("Database Topic: Financial Transaction Management with ACID Compliance")

doc.add_page_break()

# ==========================================
# II. COMPONENTS OF DATABASE DESIGN
# ==========================================
add_section_heading("II. Components of Database Design")

add_subheading("Entities and Attributes")

# User
add_sub3("1. User")
add_table(
    ["Attribute", "Data Type", "Constraint", "Description"],
    [
        ["id", "VARCHAR (CUID)", "PRIMARY KEY", "Unique identifier for each user"],
        ["username", "VARCHAR", "UNIQUE, NOT NULL", "User's display name"],
        ["email", "VARCHAR", "UNIQUE, NOT NULL", "User's email address"],
        ["createdAt", "TIMESTAMP", "DEFAULT now()", "Account creation date"],
    ]
)

# Wallet
add_sub3("2. Wallet")
add_table(
    ["Attribute", "Data Type", "Constraint", "Description"],
    [
        ["id", "VARCHAR (CUID)", "PRIMARY KEY", "Unique identifier for the wallet"],
        ["balance", "DECIMAL", "DEFAULT 10000.00", "Current cash balance in USD"],
        ["userId", "VARCHAR", "FK → User(id), UNIQUE", "Links wallet to its owner"],
    ]
)

# Asset
add_sub3("3. Asset")
add_table(
    ["Attribute", "Data Type", "Constraint", "Description"],
    [
        ["id", "VARCHAR (CUID)", "PRIMARY KEY", "Unique identifier for each asset"],
        ["symbol", "VARCHAR", "UNIQUE, NOT NULL", "Ticker symbol (e.g., AAPL, BTC)"],
        ["name", "VARCHAR", "NOT NULL", "Full name of the asset"],
        ["currentPrice", "DECIMAL", "NOT NULL", "Current live market price"],
    ]
)

# Portfolio
add_sub3("4. Portfolio")
add_table(
    ["Attribute", "Data Type", "Constraint", "Description"],
    [
        ["id", "VARCHAR (CUID)", "PRIMARY KEY", "Unique identifier for the holding"],
        ["userId", "VARCHAR", "FK → User(id)", "The user who owns this holding"],
        ["assetId", "VARCHAR", "FK → Asset(id)", "The asset being held"],
        ["quantity", "DECIMAL", "NOT NULL", "Number of units owned"],
        ["", "", "UNIQUE(userId, assetId)", "One row per user per asset"],
    ]
)

# Transaction
add_sub3("5. Transaction")
add_table(
    ["Attribute", "Data Type", "Constraint", "Description"],
    [
        ["id", "VARCHAR (CUID)", "PRIMARY KEY", "Unique identifier for the transaction"],
        ["userId", "VARCHAR", "FK → User(id)", "The user who made the trade"],
        ["assetId", "VARCHAR", "NOT NULL", "The asset that was traded"],
        ["type", "VARCHAR", "NOT NULL", "Trade direction: 'BUY' or 'SELL'"],
        ["amount", "DECIMAL", "NOT NULL", "Number of units traded"],
        ["price", "DECIMAL", "NOT NULL", "Price per unit at time of trade"],
        ["timestamp", "TIMESTAMP", "DEFAULT now()", "When the trade was executed"],
    ]
)

add_subheading("Relationships")
add_table(
    ["Relationship", "Entities", "Cardinality", "Participation"],
    [
        ["has_wallet", "User → Wallet", "1 : 1", "User: Total, Wallet: Total"],
        ["holds", "User → Portfolio", "1 : N", "User: Partial, Portfolio: Total"],
        ["held_in", "Asset → Portfolio", "1 : N", "Asset: Partial, Portfolio: Total"],
        ["makes", "User → Transaction", "1 : N", "User: Partial, Transaction: Total"],
    ]
)

add_body("Each User has exactly one Wallet (1:1). A User can hold many assets in their Portfolio (1:N). Each Asset can appear in many users' Portfolios (1:N). A User can make many Transactions (1:N). Each Transaction is always tied to a User.")

doc.add_page_break()

# ==========================================
# III. ER DIAGRAM
# ==========================================
add_section_heading("III. Entity Relationship Diagram")

add_body("The ER diagram below shows the 5 entities and their relationships:")
add_body("")

# Since we can't render mermaid in a docx, we'll describe it textually AND draw it with a table-based diagram
add_body("USER  ——(1:1)——  WALLET")
add_body("USER  ——(1:N)——  PORTFOLIO  ——(N:1)——  ASSET")
add_body("USER  ——(1:N)——  TRANSACTION")
add_body("")

add_body("Detailed relationships:")
add_bullet("USER ||--|| WALLET : 'has_wallet' (One-to-One)")
add_bullet("USER ||--o{ PORTFOLIO : 'holds' (One-to-Many)")
add_bullet("ASSET ||--o{ PORTFOLIO : 'held_in' (One-to-Many)")
add_bullet("USER ||--o{ TRANSACTION : 'makes' (One-to-Many)")

add_body("")
add_body("Note: The ER diagram can be drawn using tools like draw.io, Lucidchart, or MySQL Workbench. The textual representation above maps directly to the standard Chen/Crow's Foot notation.")

doc.add_page_break()

# ==========================================
# IV. RELATIONAL MODEL
# ==========================================
add_section_heading("IV. Relational Model")

add_body("Converting the ER diagram into relational tables:")
add_body("")

add_code("User(id, username, email, createdAt)\n    Primary Key: id\n    Unique: username, email")
add_body("")
add_code("Wallet(id, balance, userId)\n    Primary Key: id\n    Foreign Key: userId → User(id)\n    Unique: userId")
add_body("")
add_code("Asset(id, symbol, name, currentPrice)\n    Primary Key: id\n    Unique: symbol")
add_body("")
add_code("Portfolio(id, userId, assetId, quantity)\n    Primary Key: id\n    Foreign Key: userId → User(id)\n    Foreign Key: assetId → Asset(id)\n    Unique: (userId, assetId)  ← Composite unique key")
add_body("")
add_code("Transaction(id, userId, assetId, type, amount, price, timestamp)\n    Primary Key: id\n    Foreign Key: userId → User(id)")

add_body("")
add_body("Total Relations: 5")

doc.add_page_break()

# ==========================================
# V. NORMALIZATION
# ==========================================
add_section_heading("V. Normalization")

add_subheading("First Normal Form (1NF) ✓")
add_body("Rule: All attributes must contain atomic (indivisible) values. No repeating groups.")
add_body("All tables satisfy 1NF. Every attribute holds a single value (e.g., username is a single string, not a list). There are no repeating groups — a user's assets are stored in separate rows in Portfolio, not as a comma-separated list. Every table has a defined primary key.")

add_subheading("Second Normal Form (2NF) ✓")
add_body("Rule: Must be in 1NF + no partial dependency (non-key attributes must depend on the entire primary key).")
add_body("All tables use a single-column primary key (id), so partial dependency is impossible. In the Portfolio table, while there is a composite unique key (userId, assetId), the primary key is still the single column id. Therefore, all non-key attributes depend fully on id.")

add_subheading("Third Normal Form (3NF) ✓")
add_body("Rule: Must be in 2NF + no transitive dependency (non-key attributes must not depend on other non-key attributes).")
add_bullet("User: username and email depend only on id, not on each other")
add_bullet("Wallet: balance depends only on id, not on userId")
add_bullet("Asset: name and currentPrice depend only on id, not on symbol")
add_bullet("Portfolio: quantity depends only on id")
add_bullet("Transaction: type, amount, price, timestamp all depend only on id")
add_body("No transitive dependencies found.")

add_subheading("Boyce-Codd Normal Form (BCNF) ✓")
add_body("Rule: Must be in 3NF + for every functional dependency X → Y, X must be a superkey.")
add_body("In all five tables, the only determinant is the primary key id, which is always a superkey. The unique constraints on username, email, symbol, and userId (in Wallet) are candidate keys — they also act as superkeys. Therefore, all tables are in BCNF.")

doc.add_page_break()

# ==========================================
# VI. SQL QUERIES
# ==========================================
add_section_heading("VI. SQL Queries")
add_body("Database: PostgreSQL (hosted on Supabase). Below are 25 queries covering all major DBMS concepts.")

# --- DDL ---
add_subheading("DDL — Creating Tables")

add_query_block(1, "Create the User table",
'''CREATE TABLE "User" (
    id       TEXT PRIMARY KEY DEFAULT gen_random_uuid(),
    username TEXT NOT NULL UNIQUE,
    email    TEXT NOT NULL UNIQUE,
    "createdAt" TIMESTAMP DEFAULT NOW()
);''',
"Creates the User table with unique constraints on username and email.")

add_query_block(2, "Create the Wallet table with Foreign Key",
'''CREATE TABLE "Wallet" (
    id      TEXT PRIMARY KEY DEFAULT gen_random_uuid(),
    balance DECIMAL DEFAULT 10000.00,
    "userId" TEXT NOT NULL UNIQUE,
    FOREIGN KEY ("userId") REFERENCES "User"(id)
);''',
"Creates Wallet with a 1:1 relationship to User via a UNIQUE foreign key.")

add_query_block(3, "Create the Asset table",
'''CREATE TABLE "Asset" (
    id            TEXT PRIMARY KEY DEFAULT gen_random_uuid(),
    symbol        TEXT NOT NULL UNIQUE,
    name          TEXT NOT NULL,
    "currentPrice" DECIMAL NOT NULL
);''',
"Creates Asset table. The symbol is unique since each ticker can only appear once.")

add_query_block(4, "Create the Portfolio table with Composite Unique Key",
'''CREATE TABLE "Portfolio" (
    id        TEXT PRIMARY KEY DEFAULT gen_random_uuid(),
    "userId"  TEXT NOT NULL,
    "assetId" TEXT NOT NULL,
    quantity  DECIMAL NOT NULL,
    FOREIGN KEY ("userId") REFERENCES "User"(id),
    FOREIGN KEY ("assetId") REFERENCES "Asset"(id),
    UNIQUE ("userId", "assetId")
);''',
"The composite unique key (userId, assetId) ensures a user can only have one portfolio entry per asset.")

add_query_block(5, "Create the Transaction table",
'''CREATE TABLE "Transaction" (
    id        TEXT PRIMARY KEY DEFAULT gen_random_uuid(),
    "userId"  TEXT NOT NULL,
    "assetId" TEXT NOT NULL,
    type      TEXT NOT NULL,
    amount    DECIMAL NOT NULL,
    price     DECIMAL NOT NULL,
    timestamp TIMESTAMP DEFAULT NOW(),
    FOREIGN KEY ("userId") REFERENCES "User"(id)
);''',
"Records every trade. The type column stores 'BUY' or 'SELL'.")

# --- DML ---
add_subheading("DML — Inserting Data")

add_query_block(6, "Insert Users",
'''INSERT INTO "User" (id, username, email) VALUES
    ('user1', 'aditya_trade', 'aditya@test.com'),
    ('user2', 'rahul_invest', 'rahul@test.com'),
    ('user3', 'priya_stocks', 'priya@test.com');''',
"Inserts 3 sample users into the database.")

add_query_block(7, "Insert Wallets",
'''INSERT INTO "Wallet" (id, balance, "userId") VALUES
    ('w1', 10000.00, 'user1'),
    ('w2', 10000.00, 'user2'),
    ('w3', 10000.00, 'user3');''',
"Creates a wallet with $10,000 starting balance for each user.")

add_query_block(8, "Insert Market Assets",
'''INSERT INTO "Asset" (id, symbol, name, "currentPrice") VALUES
    ('a1', 'AAPL',  'Apple Inc.',     180.50),
    ('a2', 'BTC',   'Bitcoin',      65000.00),
    ('a3', 'TSLA',  'Tesla, Inc.',    175.20),
    ('a4', 'ETH',   'Ethereum',      3500.00),
    ('a5', 'GOOGL', 'Alphabet Inc.',  150.10);''',
"Populates the market with 5 tradable assets and their initial prices.")

add_query_block(9, "Insert Portfolio Holdings",
'''INSERT INTO "Portfolio" (id, "userId", "assetId", quantity) VALUES
    ('p1', 'user1', 'a1', 10),   ('p2', 'user1', 'a2', 0.05),
    ('p3', 'user2', 'a3', 15),   ('p4', 'user2', 'a4', 1.5),
    ('p5', 'user3', 'a1', 20),   ('p6', 'user3', 'a5', 25),
    ('p7', 'user1', 'a4', 2),    ('p8', 'user2', 'a1', 5),
    ('p9', 'user3', 'a2', 0.02), ('p10','user1', 'a5', 8);''',
"Inserts 10 portfolio holdings across 3 users.")

add_query_block(10, "Insert Transaction Records",
'''INSERT INTO "Transaction" (id, "userId", "assetId", type, amount, price, timestamp) VALUES
    ('t1',  'user1', 'a1', 'BUY',  10,   180.50, '2026-04-10 09:30:00'),
    ('t2',  'user1', 'a2', 'BUY',  0.05, 65000,  '2026-04-10 10:15:00'),
    ('t3',  'user2', 'a3', 'BUY',  15,   175.20, '2026-04-10 11:00:00'),
    ('t4',  'user2', 'a4', 'BUY',  1.5,  3500,   '2026-04-11 09:00:00'),
    ('t5',  'user3', 'a1', 'BUY',  20,   180.50, '2026-04-11 10:30:00'),
    ('t6',  'user3', 'a5', 'BUY',  25,   150.10, '2026-04-11 14:00:00'),
    ('t7',  'user1', 'a4', 'BUY',  2,    3500,   '2026-04-12 09:15:00'),
    ('t8',  'user1', 'a1', 'SELL', 3,    185.00, '2026-04-12 13:00:00'),
    ('t9',  'user2', 'a3', 'SELL', 5,    178.00, '2026-04-13 10:00:00'),
    ('t10', 'user3', 'a5', 'SELL', 10,   155.00, '2026-04-13 15:30:00'),
    ('t11', 'user1', 'a5', 'BUY',  8,    150.10, '2026-04-14 09:00:00'),
    ('t12', 'user2', 'a1', 'BUY',  5,    182.00, '2026-04-14 11:00:00');''',
"Inserts 12 transaction records (mix of BUY and SELL) over multiple days.")

# --- SELECT ---
add_subheading("SELECT Queries")

add_query_block(11, "Get all assets sorted by price (ORDER BY)",
'''SELECT symbol, name, "currentPrice"
FROM "Asset"
ORDER BY "currentPrice" DESC;''',
"Lists all tradable assets from most to least expensive.")

add_query_block(12, "Find high-volume traders (GROUP BY + HAVING)",
'''SELECT u.username,
       COUNT(t.id) AS total_trades,
       SUM(t.amount * t.price) AS total_volume
FROM "Transaction" t
JOIN "User" u ON t."userId" = u.id
GROUP BY u.username
HAVING SUM(t.amount * t.price) > 1000
ORDER BY total_volume DESC;''',
"Identifies traders whose total trade volume exceeds $1000 using aggregation and HAVING filter.")

# --- JOINs ---
add_subheading("JOIN Queries")

add_query_block(13, "Inner Join — Portfolio with market values",
'''SELECT u.username, a.symbol, a.name, p.quantity,
       a."currentPrice",
       (p.quantity * a."currentPrice") AS market_value
FROM "Portfolio" p
INNER JOIN "User" u ON p."userId" = u.id
INNER JOIN "Asset" a ON p."assetId" = a.id
ORDER BY market_value DESC;''',
"Shows each user's holdings with live market valuations by joining three tables.")

add_query_block(14, "Left Join — All users including those with no trades",
'''SELECT u.username, COUNT(t.id) AS trade_count
FROM "User" u
LEFT JOIN "Transaction" t ON u.id = t."userId"
GROUP BY u.username;''',
"Shows every user even if they haven't made any trades yet (trade_count = 0).")

# --- Subqueries ---
add_subheading("Subqueries")

add_query_block(15, "Users with below-average wallet balance (Scalar Subquery)",
'''SELECT u.username, w.balance
FROM "User" u
JOIN "Wallet" w ON u.id = w."userId"
WHERE w.balance < (
    SELECT AVG(balance) FROM "Wallet"
);''',
"Uses a scalar subquery to compare each user's balance against the global average.")

add_query_block(16, "Assets that no user holds (NOT IN Subquery)",
'''SELECT symbol, name
FROM "Asset"
WHERE id NOT IN (
    SELECT DISTINCT "assetId" FROM "Portfolio"
);''',
"Identifies assets with zero demand across all users using a NOT IN subquery.")

# --- Aggregates ---
add_subheading("Aggregate Functions")

add_query_block(17, "Full trading statistics per user (COUNT, SUM, AVG, MIN, MAX, CASE)",
'''SELECT u.username,
       COUNT(t.id) AS total_trades,
       SUM(CASE WHEN t.type = 'BUY' THEN 1 ELSE 0 END) AS buys,
       SUM(CASE WHEN t.type = 'SELL' THEN 1 ELSE 0 END) AS sells,
       ROUND(SUM(t.amount * t.price)::numeric, 2) AS total_volume,
       ROUND(AVG(t.amount * t.price)::numeric, 2) AS avg_trade_size,
       MIN(t.timestamp) AS first_trade,
       MAX(t.timestamp) AS last_trade
FROM "Transaction" t
JOIN "User" u ON t."userId" = u.id
GROUP BY u.username;''',
"Comprehensive statistics using COUNT, SUM, AVG, MIN, MAX, and CASE expressions.")

# --- View ---
add_subheading("View")

add_query_block(18, "Create a View for portfolio dashboard",
'''CREATE VIEW user_portfolio_dashboard AS
SELECT u.username, u.email, w.balance AS wallet_balance,
       COALESCE(SUM(p.quantity * a."currentPrice"), 0) AS portfolio_value,
       w.balance + COALESCE(SUM(p.quantity * a."currentPrice"), 0) AS net_worth
FROM "User" u
JOIN "Wallet" w ON u.id = w."userId"
LEFT JOIN "Portfolio" p ON u.id = p."userId"
LEFT JOIN "Asset" a ON p."assetId" = a.id
GROUP BY u.username, u.email, w.balance;

-- Usage:
SELECT * FROM user_portfolio_dashboard ORDER BY net_worth DESC;''',
"Creates a reusable view that always shows up-to-date portfolio metrics.")

# --- UPDATE / DELETE ---
add_subheading("UPDATE and DELETE")

add_query_block(19, "Update asset price (simulating market movement)",
'''UPDATE "Asset"
SET "currentPrice" = "currentPrice" * 1.05
WHERE symbol = 'AAPL';''',
"Increases Apple's price by 5% to simulate market volatility.")

add_query_block(20, "Delete a portfolio entry when all shares are sold",
'''DELETE FROM "Portfolio"
WHERE "userId" = 'user1'
  AND "assetId" = 'a3'
  AND quantity = 0;''',
"Cleans up portfolio entries with zero holdings.")

# --- ACID Transaction ---
add_subheading("ACID Transaction (TCL)")

add_query_block(21, "Execute a BUY trade with full ACID compliance",
'''BEGIN;

-- Step 1: Deduct cost from wallet
UPDATE "Wallet" SET balance = balance - (2 * 180.50)
WHERE "userId" = 'user1' AND balance >= (2 * 180.50);

-- Step 2: Add to portfolio (upsert)
INSERT INTO "Portfolio" (id, "userId", "assetId", quantity)
VALUES (gen_random_uuid(), 'user1', 'a1', 2)
ON CONFLICT ("userId", "assetId")
DO UPDATE SET quantity = "Portfolio".quantity + 2;

-- Step 3: Record the transaction
INSERT INTO "Transaction" (id, "userId", "assetId", type, amount, price)
VALUES (gen_random_uuid(), 'user1', 'a1', 'BUY', 2, 180.50);

COMMIT;
-- If any step fails, use ROLLBACK instead of COMMIT''',
"Full ACID transaction. All three operations succeed together or fail together. This is the core DBMS concept — ensuring data consistency during financial operations.")

# --- Window Functions ---
add_subheading("Window Functions")

add_query_block(22, "Rank users by portfolio value using RANK()",
'''SELECT username, portfolio_value,
       RANK() OVER (ORDER BY portfolio_value DESC) AS rank
FROM user_portfolio_dashboard;''',
"Uses a window function to rank users by portfolio value without collapsing rows.")

add_query_block(23, "Running total of trades using SUM() OVER",
'''SELECT u.username, t.type, a.symbol,
       t.amount * t.price AS trade_value,
       t.timestamp,
       SUM(t.amount * t.price) OVER (
           PARTITION BY t."userId" ORDER BY t.timestamp
       ) AS running_total
FROM "Transaction" t
JOIN "User" u ON t."userId" = u.id
JOIN "Asset" a ON t."assetId" = a.id
ORDER BY t."userId", t.timestamp;''',
"Calculates a cumulative running total per user using PARTITION BY.")

# --- Trigger ---
add_subheading("Trigger and PL/pgSQL Function")

add_query_block(24, "Trigger to auto-log wallet balance changes",
'''-- Step 1: Create audit log table
CREATE TABLE wallet_audit_log (
    id SERIAL PRIMARY KEY, wallet_id TEXT,
    old_balance DECIMAL, new_balance DECIMAL,
    changed_at TIMESTAMP DEFAULT NOW()
);

-- Step 2: Create the trigger function
CREATE OR REPLACE FUNCTION log_wallet_change()
RETURNS TRIGGER AS $$
BEGIN
    INSERT INTO wallet_audit_log (wallet_id, old_balance, new_balance)
    VALUES (OLD.id, OLD.balance, NEW.balance);
    RETURN NEW;
END;
$$ LANGUAGE plpgsql;

-- Step 3: Attach the trigger
CREATE TRIGGER wallet_balance_trigger
AFTER UPDATE OF balance ON "Wallet"
FOR EACH ROW EXECUTE FUNCTION log_wallet_change();''',
"Automatically records every wallet balance change. Demonstrates triggers and PL/pgSQL.")

# --- Stored Procedure ---
add_subheading("Stored Procedure")

add_query_block(25, "Get full user summary",
'''CREATE OR REPLACE FUNCTION get_user_summary(user_email TEXT)
RETURNS TABLE (
    username TEXT, wallet_balance DECIMAL,
    total_holdings BIGINT, total_trades BIGINT,
    portfolio_value DECIMAL
) AS $$
BEGIN
    RETURN QUERY
    SELECT u.username, w.balance,
           COUNT(DISTINCT p.id), COUNT(DISTINCT t.id),
           COALESCE(SUM(p.quantity * a."currentPrice"), 0)
    FROM "User" u
    JOIN "Wallet" w ON u.id = w."userId"
    LEFT JOIN "Portfolio" p ON u.id = p."userId"
    LEFT JOIN "Asset" a ON p."assetId" = a.id
    LEFT JOIN "Transaction" t ON u.id = t."userId"
    WHERE u.email = user_email
    GROUP BY u.username, w.balance;
END;
$$ LANGUAGE plpgsql;

-- Usage:
SELECT * FROM get_user_summary('aditya@test.com');''',
"A reusable stored procedure that returns a complete user summary in one call.")

# --- Concepts Summary Table ---
add_subheading("Summary of DBMS Concepts Covered")
add_table(
    ["#", "Concept", "Query #"],
    [
        ["1", "DDL (CREATE TABLE)", "1–5"],
        ["2", "DML (INSERT)", "6–10"],
        ["3", "SELECT, ORDER BY", "11"],
        ["4", "GROUP BY, HAVING", "12"],
        ["5", "INNER JOIN (multi-table)", "13"],
        ["6", "LEFT JOIN", "14"],
        ["7", "Scalar Subquery", "15"],
        ["8", "NOT IN Subquery", "16"],
        ["9", "Aggregate Functions", "17"],
        ["10", "CASE Expression", "17"],
        ["11", "CREATE VIEW", "18"],
        ["12", "UPDATE", "19"],
        ["13", "DELETE", "20"],
        ["14", "ACID Transaction (BEGIN/COMMIT)", "21"],
        ["15", "UPSERT (ON CONFLICT)", "21"],
        ["16", "Window Function (RANK)", "22"],
        ["17", "Window Function (SUM OVER)", "23"],
        ["18", "Trigger", "24"],
        ["19", "PL/pgSQL Function", "24, 25"],
        ["20", "Stored Procedure", "25"],
    ]
)

doc.add_page_break()

# ==========================================
# VII. PROJECT DEMONSTRATION
# ==========================================
add_section_heading("VII. Project Demonstration")

add_subheading("Tools / Software / Libraries Used")
add_table(
    ["Category", "Tool", "Purpose"],
    [
        ["Database", "PostgreSQL (Supabase)", "Primary RDBMS"],
        ["ORM", "Prisma 7", "Type-safe database access from Node.js"],
        ["Backend", "Express.js (Node.js)", "REST API server"],
        ["Frontend", "Angular 21 (TypeScript)", "Single-page application UI"],
        ["Charts", "Chart.js", "Live price graphs and portfolio charts"],
        ["Styling", "Tailwind CSS 3", "Utility-first CSS framework"],
        ["Market Sim", "Custom Node.js script", "Updates asset prices every 3 seconds"],
        ["Hosting", "Supabase (cloud PostgreSQL)", "Database hosting with SSL"],
    ]
)

add_subheading("Application Pages")

add_body("1. Dashboard — Shows wallet balance, portfolio value, net worth, and a real-time multi-asset price chart with asset tabs (AAPL, BTC, ETH, GOOGL, TSLA). Includes a live market table with session change indicators.")

add_body("2. Trade — Users can select an asset, choose BUY or SELL, enter a quantity, and see a live cost calculator. A confirmation modal shows the order summary. The trade uses a database transaction (ACID).")

add_body("3. Portfolio — Displays all current holdings with quantities, current price, market value, and allocation percentage. A doughnut chart visualizes the portfolio breakdown.")

add_body("4. History — Shows a complete log of all past transactions with date/time, trade type (BUY/SELL), asset symbol, quantity, price, and total value. Supports filtering by asset and type.")

doc.add_page_break()

# ==========================================
# VIII. SELF-LEARNING
# ==========================================
add_section_heading("VIII. Self-Learning Beyond Classroom")

add_bullet("Prisma ORM Interactive Transactions — Learned how to wrap multiple database operations inside prisma.$transaction() to ensure atomicity, which maps directly to SQL's BEGIN/COMMIT/ROLLBACK")
add_bullet("UPSERT Operations — Learned PostgreSQL's ON CONFLICT ... DO UPDATE syntax, used when a user buys more of an asset they already hold")
add_bullet("Connection Pooling — Learned how database connection pools (via the pg library) improve performance by reusing connections instead of opening a new one for each query")
add_bullet("Real-time Data Architecture — Learned how to build a market simulation service that updates prices using setInterval and how Angular polls the API every 3 seconds for live updates")
add_bullet("Cloud Database Management — Learned to configure Supabase (cloud-hosted PostgreSQL) with SSL, environment variables, and connection poolers")

# ==========================================
# IX. LEARNING FROM PROJECT
# ==========================================
add_section_heading("IX. Learning from the Project")

add_bullet("ACID transactions are essential for financial applications — Without wrapping BUY trades in a transaction, money could be deducted from the wallet but the asset not added to the portfolio, leading to data inconsistency.")
add_bullet("Database schema design directly impacts application complexity — By using a composite unique key (userId, assetId) on Portfolio, we simplified the buy logic to a simple UPSERT.")
add_bullet("Normalization prevents data anomalies — Storing asset prices only in the Asset table means price updates happen in one place and are reflected everywhere.")
add_bullet("Foreign keys enforce data integrity — Even if the application code has bugs, the database prevents orphaned records.")
add_bullet("Views simplify complex queries — The user_portfolio_dashboard view encapsulates a 4-table join, making it reusable without repeating complex SQL.")

# ==========================================
# X. CHALLENGES
# ==========================================
add_section_heading("X. Challenges Faced")

add_body("1. ACID Transaction Design — Designing the trade execution to be truly atomic was challenging. A single BUY operation requires 3 coordinated database operations (deduct wallet, update portfolio, insert transaction). We used Prisma's interactive transactions to wrap all three.")

add_body("2. Handling Concurrent Price Updates — The market simulator updates prices every 3 seconds. If a user initiates a trade at the same moment, the trade could use stale prices. We solved this by reading the asset price inside the transaction.")

add_body("3. Composite Unique Constraints — Implementing 'one portfolio entry per user per asset' required a composite unique key and UPSERT logic. Initial attempts with separate INSERT/UPDATE caused race conditions.")

add_body("4. Decimal Precision — Financial applications require precise decimal arithmetic. JavaScript's floating-point math was problematic. We used PostgreSQL's DECIMAL type and Prisma's Decimal to ensure accuracy.")

add_body("5. Cross-Origin Communication — The Angular frontend (port 4200) needed to communicate with the Express backend (port 3000). We configured CORS middleware and learned about the Same-Origin Policy.")

# ==========================================
# XI. CONCLUSION
# ==========================================
add_section_heading("XI. Conclusion")

add_body("The VirtualTrade project successfully demonstrates how DBMS concepts apply to a real-world financial application:")

add_bullet("ACID transactions ensure that every trade either fully succeeds or fully fails — money is never lost or created out of thin air")
add_bullet("Referential integrity (foreign keys) guarantees that portfolio entries and transactions always point to valid users and assets")
add_bullet("Normalization (BCNF) eliminates data redundancy — asset prices are stored in one place and reflected everywhere")
add_bullet("Views, triggers, and stored procedures extend the database with reusable logic beyond basic CRUD")
add_bullet("Window functions enable analytics like user rankings and running totals without collapsing data")

add_body("")
add_body("The project bridges the gap between theoretical DBMS concepts learned in class and practical application development. Building a live trading simulator made abstract concepts like atomicity, consistency, and referential integrity tangible and easy to understand.")

add_body("")
p_final = doc.add_paragraph()
run_final = p_final.add_run("Key Takeaway: A well-designed database schema with proper constraints, relationships, and transactions is the foundation of any reliable software system — the frontend is just a window into the database.")
run_final.font.bold = True
run_final.font.italic = True

# ==========================================
# SAVE
# ==========================================
output_path = "/Users/adixya7769/Documents/virtual-trade/DBMS_Report_VirtualTrade.docx"
doc.save(output_path)
print(f"✅ Report saved to: {output_path}")
